import sys
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

sys.stdout.reconfigure(encoding='utf-8')

from create_essay_helpers import (
    set_cell_background, set_cell_margins,
    add_heading_1, add_heading_2, add_heading_3,
    add_paragraph, add_bullet, add_image_placeholder, add_code_block
)

doc = docx.Document()

# Page setup: A4, Margins: Top=2cm, Bottom=2cm, Left=3cm, Right=2cm
sections = doc.sections
for section in sections:
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.79)
    section.bottom_margin = Inches(0.79)
    section.left_margin = Inches(1.18)
    section.right_margin = Inches(0.79)

# --------------------------------------------------------------------
# 1. TRANG BÌA CHÍNH (COVER PAGE)
# --------------------------------------------------------------------
p_school = doc.add_paragraph()
p_school.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_school.add_run("TRƯỜNG ĐẠI HỌC NGUYỄN TẤT THÀNH\nKHOA CÔNG NGHỆ THÔNG TIN")
r.font.name = 'Times New Roman'
r.font.size = Pt(14)
r.font.bold = True
r.font.color.rgb = RGBColor(195, 45, 33)

p_space1 = doc.add_paragraph()
p_space1.paragraph_format.space_before = Pt(36)

p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_t1 = p_title.add_run("BÁO CÁO TIỂU LUẬN / ĐỒ ÁN MÔN HỌC\n\n")
r_t1.font.name = 'Times New Roman'
r_t1.font.size = Pt(16)
r_t1.font.bold = True
r_t1.font.color.rgb = RGBColor(30, 40, 60)

r_t2 = p_title.add_run("ĐỀ TÀI:\nXÂY DỰNG ỨNG DỤNG DI ĐỘNG QUẢN LÝ HỌC TẬP TRỰC TUYẾN VIETSYNCMOBILE TÍCH HỢP SUPABASE BACKEND-AS-A-SERVICE VÀ KIẾN TRÚC MVVM TRÊN HỆ ĐIỀU HÀNH ANDROID")
r_t2.font.name = 'Times New Roman'
r_t2.font.size = Pt(18)
r_t2.font.bold = True
r_t2.font.color.rgb = RGBColor(195, 45, 33)

p_space2 = doc.add_paragraph()
p_space2.paragraph_format.space_before = Pt(48)

p_info = doc.add_paragraph()
p_info.alignment = WD_ALIGN_PARAGRAPH.LEFT
p_info.paragraph_format.left_indent = Inches(1.5)
p_info.paragraph_format.line_spacing = 1.3

r_info = p_info.add_run(
    "Giảng viên hướng dẫn:  ThS. Nguyễn Văn A\n"
    "Môn học:              Phát triển Ứng dụng Di động / Đồ án Chuyên ngành\n"
    "Sinh viên thực hiện:   Đặng Tuấn - MSSV: 210000xxxx\n"
    "Lớp học phần:         21DTHxx\n"
    "Khóa học:             2021 - 2025"
)
r_info.font.name = 'Times New Roman'
r_info.font.size = Pt(13)

p_space3 = doc.add_paragraph()
p_space3.paragraph_format.space_before = Pt(48)

p_foot = doc.add_paragraph()
p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_foot = p_foot.add_run("TP. HỒ CHÍ MINH, NĂM 2026")
r_foot.font.name = 'Times New Roman'
r_foot.font.size = Pt(13)
r_foot.font.bold = True

doc.add_page_break()

# --------------------------------------------------------------------
# 2. TỜ NHIỆM VỤ TIỂU LUẬN / ĐỒ ÁN
# --------------------------------------------------------------------
p_task_title = doc.add_paragraph()
p_task_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_tt = p_task_title.add_run("TRƯỜNG ĐẠI HỌC NGUYỄN TẤT THÀNH\nKHOA CÔNG NGHỆ THÔNG TIN\n\nNHIỆM VỤ TIỂU LUẬN / ĐỒ ÁN MÔN HỌC")
r_tt.font.name = 'Times New Roman'
r_tt.font.size = Pt(14)
r_tt.font.bold = True

add_paragraph(doc, "Họ và tên sinh viên: Đặng Tuấn", bold_prefix="1. ")
add_paragraph(doc, "Mã số sinh viên: 210000xxxx - Lớp: 21DTHxx", bold_prefix="2. ")
add_paragraph(doc, "Xây dựng ứng dụng di động Quản lý học tập trực tuyến VietsyncMobile tích hợp Supabase Backend-as-a-Service và Kiến trúc MVVM trên hệ điều hành Android.", bold_prefix="3. Tên đề tài: ")
add_paragraph(doc, "Từ ngày 01/01/2026 đến ngày 30/07/2026", bold_prefix="4. Thời gian thực hiện: ")
add_paragraph(doc, "Nhiệm vụ được giao:", bold_prefix="5. ")
add_bullet(doc, "Nghiên cứu kiến trúc MVVM, Repository Pattern và Supabase Backend-as-a-Service.")
add_bullet(doc, "Thiết kế và cài đặt cơ sở dữ liệu PostgreSQL trên Supabase RLS và Room DB local.")
add_bullet(doc, "Phát triển các module chức năng cốt lõi: Xác thực tài khoản (JWT, Refresh Token, Password Recovery), Thi cuối khóa (Final Exam Timer 30p, Navigation Grid), Cấp chứng chỉ PDF kèm QR Code xác thực, Sổ điểm Analytics, Lịch sử học tập Heatmap 90 ngày, Đánh giá ngang hàng Peer Review, Hệ thống Huy hiệu Gamification, Dashboard Giảng viên.")
add_bullet(doc, "Tích hợp Firebase Cloud Messaging (FCM), WorkManager background jobs, Dark Mode, Đa ngôn ngữ (VI/EN).")
add_bullet(doc, "Xây dựng bộ kiểm thử tự động Unit Testing (JUnit4, Mockito, MockWebServer, Robolectric) và UI Testing (Espresso).")
add_bullet(doc, "Cấu hình ProGuard R8 Full Mode mã hóa obfuscation và tối ưu hóa hiệu năng Android Profiler (Glide, OkHttp 50MB Cache, AppInitializer).")

p_sig = doc.add_paragraph()
p_sig.paragraph_format.space_before = Pt(24)
p_sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r_sig = p_sig.add_run("TP. Hồ Chí Minh, ngày 30 tháng 07 năm 2026\nGiảng viên hướng dẫn\n\n\n\nThS. Nguyễn Văn A")
r_sig.font.name = 'Times New Roman'
r_sig.font.size = Pt(13)
r_sig.font.italic = True

doc.add_page_break()

# --------------------------------------------------------------------
# 3. LỜI CẢM ƠN
# --------------------------------------------------------------------
p_thanks = doc.add_paragraph()
p_thanks.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_th = p_thanks.add_run("LỜI CẢM ƠN")
r_th.font.name = 'Times New Roman'
r_th.font.size = Pt(14)
r_th.font.bold = True

add_paragraph(doc, "Trong quá trình học tập, nghiên cứu và hoàn thành đề tài tiểu luận môn học này, em đã nhận được sự hỗ trợ, hướng dẫn tận tình và động viên chân thành từ quý Thầy Cô, gia đình cùng bạn bè.")
add_paragraph(doc, "Trước hết, em xin gửi lời cảm ơn sâu sắc đến Ban Giám hiệu Trường Đại học Nguyễn Tất Thành, quý Thầy Cô Khoa Công nghệ Thông tin đã tạo điều kiện học tập tốt nhất và truyền đạt những kiến thức chuyên môn vô cùng quý báu trong suốt thời gian qua.")
add_paragraph(doc, "Đặc biệt, em xin bày tỏ lòng biết ơn chân thành và sâu sắc nhất tới ThS. Nguyễn Văn A - Giảng viên hướng dẫn trực tiếp môn học. Thầy đã dành nhiều thời gian tận tụy chỉ bảo, định hướng phương pháp nghiên cứu và đưa ra những lời khuyên chuyên môn sắc bén giúp em từng bước hoàn thiện đề tài VietsyncMobile một cách trọn vẹn nhất.")
add_paragraph(doc, "Mặc dù đã cố gắng hết sức để hoàn thiện dự án theo đúng chuẩn mực khoa học và công nghệ hiện đại, song bài tiểu luận chắc chắn khó tránh khỏi những hạn chế nhất định. Em rất mong nhận được những ý kiến đóng góp, nhận xét và phản hồi quý báu từ quý Thầy Cô để dự án ngày càng hoàn thiện hơn.")
add_paragraph(doc, "Em xin chân thành cảm ơn!", italic=True)

doc.add_page_break()

# --------------------------------------------------------------------
# 4. PHIẾU CHẤM THI TIỂU LUẬN / ĐỒ ÁN
# --------------------------------------------------------------------
p_rub = doc.add_paragraph()
p_rub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_rb = p_rub.add_run("PHIẾU CHẤM THI TIỂU LUẬN / ĐỒ ÁN MÔN HỌC")
r_rb.font.name = 'Times New Roman'
r_rb.font.size = Pt(14)
r_rb.font.bold = True

add_paragraph(doc, "Môn thi: Phát triển Ứng dụng Di động / Đồ án Chuyên ngành      Lớp học phần: 21DTHxx")
add_paragraph(doc, "Nhóm sinh viên thực hiện: Đặng Tuấn - MSSV: 210000xxxx (Tham gia đóng góp: 100%)")
add_paragraph(doc, "Đề tài tiểu luận: Xây dựng ứng dụng di động Quản lý học tập trực tuyến VietsyncMobile tích hợp Supabase BaaS và Kiến trúc MVVM trên Android.")

# Table rubric
table = doc.add_table(rows=7, cols=4)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False

headers = ["Tiêu chí (theo CĐR HP)", "Đánh giá của Giảng viên", "Điểm tối đa", "Điểm đạt được"]
widths = [Inches(2.5), Inches(3.2), Inches(1.0), Inches(1.0)]

for i, h in enumerate(headers):
    cell = table.cell(0, i)
    cell.width = widths[i]
    set_cell_background(cell, "C32D21")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(h)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(11)
    r.font.bold = True
    r.font.color.rgb = RGBColor(255, 255, 255)

data_rubric = [
    ("Cấu trúc của báo cáo", "Đầy đủ bìa, nhiệm vụ, lời cảm ơn, mục lục, bảng biểu, chương mục chuẩn quy định Khoa CNTT", "1.0", ""),
    ("Nội dung & Đóng góp", "Phân tích yêu cầu rõ ràng, mô hình MVVM & Repository chuẩn hóa, tích hợp Supabase BaaS xuất sắc", "3.0", ""),
    ("Các nội dung thành thành", "Cài đặt hoàn chỉnh các phân hệ 6.1-7.7 (Thi cuối khóa, Chứng chỉ PDF/QR, Sổ điểm, FCM, WorkManager)", "3.0", ""),
    ("Lập luận & Kiểm thử", "Bộ Unit Test (JUnit, Mockito, Robolectric) & UI Test (Espresso) đạt 100% pass, ProGuard R8 thành công", "1.5", ""),
    ("Kết luận & Hướng phát triển", "Đánh giá ưu nhược điểm khách quan, đề xuất hướng phát triển mở rộng thực tế", "0.5", ""),
    ("Trình bày & Mã nguồn", "Định dạng văn bản đẹp, mã nguồn chuẩn hóa Clean Code, chú thích hình ảnh rõ ràng", "1.0", ""),
]

for row_idx, row_data in enumerate(data_rubric, start=1):
    for col_idx, text in enumerate(row_data):
        cell = table.cell(row_idx, col_idx)
        cell.width = widths[col_idx]
        set_cell_background(cell, "F9FAFB" if row_idx % 2 == 0 else "FFFFFF")
        set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
        p = cell.paragraphs[0]
        if col_idx in [2, 3]:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(text)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(11)

p_total = doc.add_paragraph()
p_total.paragraph_format.space_before = Pt(12)
r_tot = p_total.add_run("TỔNG ĐIỂM: ......... / 10.0  (Bằng chữ: .................................................................................)")
r_tot.font.name = 'Times New Roman'
r_tot.font.size = Pt(12)
r_tot.font.bold = True

p_gv = doc.add_paragraph()
p_gv.paragraph_format.space_before = Pt(18)
p_gv.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r_gv = p_gv.add_run("Giảng viên chấm thi\n(Ký, ghi rõ họ tên)\n\n\n\nThS. Nguyễn Văn A")
r_gv.font.name = 'Times New Roman'
r_gv.font.size = Pt(12)
r_gv.font.bold = True

doc.add_page_break()

# --------------------------------------------------------------------
# 5. MỤC LỤC & DANH MỤC
# --------------------------------------------------------------------
p_toc = doc.add_paragraph()
p_toc.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_tc = p_toc.add_run("MỤC LỤC")
r_tc.font.name = 'Times New Roman'
r_tc.font.size = Pt(14)
r_tc.font.bold = True

toc_items = [
    ("LỜI MỞ ĐẦU", "1"),
    ("CHƯƠNG 1: TỔNG QUAN ĐỀ TÀI", "2"),
    ("  1.1. Tính cấp thiết của đề tài", "2"),
    ("  1.2. Mục tiêu đề tài", "3"),
    ("  1.3. Đối tượng và Phạm vi nghiên cứu", "4"),
    ("  1.4. Phương pháp nghiên cứu và Công cụ sử dụng", "4"),
    ("  1.5. Đóng góp của đề tài", "5"),
    ("CHƯƠNG 2: CƠ SỞ LÝ THUYẾT VÀ KIẾN TRÚC HỆ THỐNG", "6"),
    ("  2.1. Tổng quan nền tảng Android Native & Ngôn ngữ Java 11", "6"),
    ("  2.2. Backend-as-a-Service (BaaS) với Supabase PostgreSQL", "7"),
    ("  2.3. Mô hình Kiến trúc phần mềm MVVM & Repository Pattern", "8"),
    ("  2.4. Công nghệ RESTful API, Retrofit 2 & OkHttp 3 Caching", "10"),
    ("  2.5. Cơ sở dữ liệu nội tuyến SQLite & Room DB", "11"),
    ("  2.6. Hệ thống Thông báo đẩy Firebase Cloud Messaging (FCM)", "12"),
    ("  2.7. Tối ưu hóa hiệu năng & Bảo mật với ProGuard R8 Full Mode", "13"),
    ("CHƯƠNG 3: MÔ HÌNH THỰC NGHIỆM VÀ CÀI ĐẶT ỨNG DỤNG", "14"),
    ("  3.1. Phân tích Yêu cầu Hệ thống", "14"),
    ("  3.2. Thiết kế Cơ sở Dữ liệu (Supabase Schema & Room DB Schema)", "15"),
    ("  3.3. Thiết kế Sơ đồ Hệ thống (Use Case, Activity, Sequence Diagrams)", "19"),
    ("  3.4. Cài đặt và Hiện thực hóa các Phân hệ Chức năng Cốt lõi", "22"),
    ("    3.4.1. Phân hệ Xác thực & Quản lý Người dùng (Auto Token Refresh, Deep-Link)", "22"),
    ("    3.4.2. Phân hệ Thi cuối khóa (Final Exam Timer 30p, Navigation Grid)", "24"),
    ("    3.4.3. Phân hệ Cấp & Xác thực Chứng chỉ Số (PDF Generator, QR Code)", "26"),
    ("    3.4.4. Phân hệ Sổ điểm & Phân tích Học tập (Grade Book, Heatmap 90 ngày)", "28"),
    ("    3.4.5. Phân hệ Peer Review (Nộp bài & Đánh giá ngang hàng ẩn danh)", "30"),
    ("    3.4.6. Phân hệ Hệ thống Huy hiệu & Thành tích (Gamification Badges)", "32"),
    ("    3.4.7. Phân hệ Dashboard Phân tích cho Giảng viên (Teacher Analytics)", "34"),
    ("    3.4.8. Phân hệ FCM Push Notifications, Dark Mode, Đa ngôn ngữ & WorkManager", "36"),
    ("  3.5. Kết quả Kiểm thử & Tối ưu hóa (Unit Test, Espresso UI Test, Release R8 APK)", "38"),
    ("CHƯƠNG 4: KẾT LUẬN VÀ HƯỚNG PHÁT TRUYỂN", "41"),
    ("  4.1. Kết luận những kết quả đạt được", "41"),
    ("  4.2. Đánh giá ưu điểm và hạn chế của ứng dụng", "42"),
    ("  4.3. Hướng phát triển trong tương lai", "43"),
    ("TÀI LIỆU THAM KHẢO", "44"),
    ("PHỤ LỤC", "45")
]

for title, page in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(title)
    r1.font.name = 'Times New Roman'
    r1.font.size = Pt(12)
    if title.startswith("CHƯƠNG") or title in ["LỜI MỞ ĐẦU", "TÀI LIỆU THAM KHẢO", "PHỤ LỤC"]:
        r1.font.bold = True
    
    dots_count = max(5, 80 - len(title))
    r_dots = p.add_run(" " + "." * dots_count + " " + page)
    r_dots.font.name = 'Times New Roman'
    r_dots.font.size = Pt(12)

doc.add_page_break()

# DANH MỤC TỪ VIẾT TẮT
p_abb = doc.add_paragraph()
p_abb.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_ab = p_abb.add_run("DANH MỤC TỪ VIẾT TẮT")
r_ab.font.name = 'Times New Roman'
r_ab.font.size = Pt(14)
r_ab.font.bold = True

abbreviations = [
    ("API", "Application Programming Interface (Giao diện lập trình ứng dụng)"),
    ("BaaS", "Backend-as-a-Service (Dịch vụ hỗ trợ Backend)"),
    ("CSV", "Comma-Separated Values (Định dạng tập tin văn bản phân tách bằng dấu phẩy)"),
    ("DAO", "Data Access Object (Đối tượng truy xuất dữ liệu local)"),
    ("FCM", "Firebase Cloud Messaging (Dịch vụ thông báo đẩy của Google)"),
    ("GPA", "Grade Point Average (Điểm trung bình tích lũy)"),
    ("JWT", "JSON Web Token (Mã xác thực cấu trúc JSON)"),
    ("LMS", "Learning Management System (Hệ thống quản lý học tập)"),
    ("MVVM", "Model-View-ViewModel (Kiến trúc phân tách View và Business Logic)"),
    ("PDF", "Portable Document Format (Định dạng tài liệu di động)"),
    ("QR", "Quick Response Code (Mã phản hồi nhanh 2D)"),
    ("RLS", "Row Level Security (Bảo mật cấp dòng trên PostgreSQL)"),
    ("SDK", "Software Development Kit (Bộ công cụ phát triển phần mềm)"),
    ("UI / UX", "User Interface / User Experience (Giao diện / Trải nghiệm người dùng)"),
]

table_abb = doc.add_table(rows=len(abbreviations)+1, cols=2)
table_abb.alignment = WD_TABLE_ALIGNMENT.CENTER
table_abb.cell(0, 0).width = Inches(1.8)
table_abb.cell(0, 1).width = Inches(4.5)

set_cell_background(table_abb.cell(0, 0), "C32D21")
set_cell_background(table_abb.cell(0, 1), "C32D21")

p0 = table_abb.cell(0, 0).paragraphs[0]
p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
r0 = p0.add_run("Từ viết tắt")
r0.font.bold = True
r0.font.color.rgb = RGBColor(255, 255, 255)

p1 = table_abb.cell(0, 1).paragraphs[0]
p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
r1 = p1.add_run("Ý nghĩa / Thuật ngữ tiếng Anh")
r1.font.bold = True
r1.font.color.rgb = RGBColor(255, 255, 255)

for idx, (short, full) in enumerate(abbreviations, start=1):
    c0 = table_abb.cell(idx, 0)
    c1 = table_abb.cell(idx, 1)
    set_cell_background(c0, "F9FAFB" if idx % 2 == 0 else "FFFFFF")
    set_cell_background(c1, "F9FAFB" if idx % 2 == 0 else "FFFFFF")
    
    p_s = c0.paragraphs[0]
    p_s.add_run(short).font.bold = True
    
    p_f = c1.paragraphs[0]
    p_f.add_run(full)

doc.add_page_break()

# --------------------------------------------------------------------
# LỜI MỞ ĐẦU
# --------------------------------------------------------------------
add_heading_1(doc, "LỜI MỞ ĐẦU")
add_paragraph(doc, "Trong kỷ nguyên chuyển đổi số bùng nổ hiện nay, giáo dục trực tuyến (E-Learning) và các Hệ thống Quản lý Học tập (Learning Management System - LMS) đã trở thành một thành phần hạ tầng không thể thiếu đối với các trường đại học, tổ chức giáo dục và doanh nghiệp. Sự phát triển mạnh mẽ của thiết bị di động thông minh chạy hệ điều hành Android đã mở ra cơ hội to lớn cho việc học tập mọi lúc, mọi nơi (M-Learning).")
add_paragraph(doc, "Tuy nhiên, hầu hết các ứng dụng LMS hiện nay trên thị trường thường gặp phải những thách thức lớn về khả năng mở rộng, chi phí vận hành hạ tầng server phức tạp, độ trễ mạng khi truy vấn dữ liệu học tập lớn, cũng như thiếu sót các tính năng tương tác chuyên sâu như thi trực tuyến ngẫu nhiên chống gian lận, cấp chứng chỉ số bảo mật bằng mã QR Code, đánh giá ngang hàng (Peer Review) và các yếu tố gamification thúc đẩy động lực học tập.")
add_paragraph(doc, "Xuất phát từ thực tiễn đó, đề tài \"Xây dựng ứng dụng di động Quản lý học tập trực tuyến VietsyncMobile tích hợp Supabase Backend-as-a-Service và Kiến trúc MVVM trên hệ điều hành Android\" được nghiên cứu và thực hiện nhằm giải quyết triệt để các vấn đề trên. Ứng dụng tận dụng sức mạnh của Supabase (nền tảng BaaS mã nguồn mở dựa trên PostgreSQL với cơ chế Row Level Security - RLS bảo mật mạnh mẽ) kết hợp với mô hình kiến trúc phần mềm chuẩn hóa MVVM (Model-View-ViewModel), Repository Pattern và các công nghệ tối ưu hiệu năng tiên tiến trên Android Native Java.")

doc.add_page_break()

# --------------------------------------------------------------------
# CHƯƠNG 1: TỔNG QUAN ĐỀ TÀI
# --------------------------------------------------------------------
add_heading_1(doc, "CHƯƠNG 1: TỔNG QUAN ĐỀ TÀI")

add_heading_2(doc, "1.1. Tính cấp thiết của đề tài")
add_paragraph(doc, "Nhu cầu học tập trực tuyến cá nhân hóa ngày càng tăng cao đòi hỏi một ứng dụng LMS di động không chỉ dừng lại ở việc xem video bài giảng tĩnh mà phải là một hệ sinh thái học tập tương tác toàn diện. Học viên cần theo dõi chi tiết tiến trình học cá nhân qua biểu đồ trực quan, thực hiện các bài thi đánh giá năng lực có giới hạn thời gian nghiêm ngặt, nhận chứng chỉ số có giá trị xác minh toàn vẹn và tham gia vào cộng đồng học tập thông qua cơ chế nộp bài và chấm điểm ngang hàng (Peer Review).")
add_paragraph(doc, "Về mặt kỹ thuật, việc tự xây dựng toàn bộ hệ thống Backend truyền thống (Node.js, Java Spring Boot, Laravel...) đòi hỏi chi phí lớn về thời gian phát triển, bảo trì máy chủ và cấu hình bảo mật. Do đó, việc ứng dụng Backend-as-a-Service (BaaS) Supabase cung cấp sẵn PostgreSQL Database, Realtime Subscriptions, Authentication JWT, và Storage Services là giải pháp kiến trúc tối ưu giúp tăng tốc độ phát triển ứng dụng lên gấp 3-4 lần mà vẫn đảm bảo tính an toàn bảo mật cấp doanh nghiệp.")

add_heading_2(doc, "1.2. Mục tiêu đề tài")
add_heading_3(doc, "1.2.1. Mục tiêu tổng quát")
add_paragraph(doc, "Xây dựng hoàn chỉnh ứng dụng di động LMS VietsyncMobile chạy trên nền tảng Android Native Java, tích hợp Supabase BaaS và áp dụng kiến trúc chuẩn mực MVVM, phục vụ đầy đủ nhu cầu của Học viên (Student) và Giảng viên (Instructor).")

add_heading_3(doc, "1.2.2. Mục tiêu cụ thể")
add_bullet(doc, "Nghiên cứu và làm chủ mô hình kiến trúc MVVM, LiveData, ViewModel, Repository Pattern và giao tiếp RESTful API qua Retrofit 2 / OkHttp 3.")
add_bullet(doc, "Thiết kế và cài đặt Cơ sở dữ liệu PostgreSQL trên Supabase với 10+ bảng dữ liệu chuẩn hóa, áp dụng mã hóa Row Level Security (RLS) và Triggers tự động.")
add_bullet(doc, "Phát triển Phân hệ Thi cuối khóa (Final Assessment) có bộ đếm thời gian 30 phút, trộn câu hỏi ngẫu nhiên, lưới điều hướng 30 câu hỏi và phân tích kết quả theo chuyên đề.")
add_bullet(doc, "Xây dựng Phân hệ Cấp chứng chỉ số tự động (PDF Generator) tích hợp mã QR Code chứa mã hash xác thực tính toàn vẹn.")
add_bullet(doc, "Phát triển Phân hệ Sổ điểm (Grade Book) và Lịch sử học tập 90 ngày dạng Heatmap Calendar giống GitHub.")
add_bullet(doc, "Xây dựng Phân hệ Peer Review đánh giá ngang hàng bài nộp ẩn danh và Phân hệ Gamification Huy hiệu (Achievement Badge System).")
add_bullet(doc, "Tích hợp Dashboard phân tích cho Giảng viên (Teacher Analytics) với đồ thị trực quan MPAndroidChart và chức năng xuất báo cáo PDF.")
add_bullet(doc, "Cài đặt hệ thống dịch vụ chạy ngầm: FCM Push Notifications, WorkManager định kỳ, Dark Mode, Đa ngôn ngữ (VI/EN).")
add_bullet(doc, "Thực hiện kiểm thử tự động Unit Testing (JUnit4, Mockito, MockWebServer, Robolectric), UI Testing (Espresso) và tối ưu hóa đóng gói mã nguồn ProGuard R8 Full Mode.")

add_heading_2(doc, "1.3. Đối tượng và Phạm vi nghiên cứu")
add_paragraph(doc, "Đối tượng nghiên cứu: Kiến trúc ứng dụng di động Android Native, mô hình MVVM, dịch vụ BaaS Supabase PostgreSQL, các công nghệ kiểm thử tự động và tối ưu hóa hiệu năng R8.")
add_paragraph(doc, "Phạm vi nghiên cứu: Ứng dụng di động VietsyncMobile chạy trên Android SDK level 24 đến 34 (Android 7.0 trở lên). Hệ thống tập trung vào các quy trình học tập trực tuyến, kiểm tra đánh giá, cấp chứng nhận và phân tích dữ liệu học tập.")

add_heading_2(doc, "1.4. Phương pháp nghiên cứu và Công cụ sử dụng")
add_paragraph(doc, "Phương pháp nghiên cứu: Kết hợp giữa nghiên cứu lý thuyết chuyên sâu (Android Developer Documentation, Supabase Docs, MVVM Guidelines) và phương pháp thực nghiệm phát triển phần mềm Agile/Scrum.")
add_paragraph(doc, "Công cụ phát triển: Android Studio Ladybug/Jellyfish, Java JDK 11, Postman, Git/GitHub, Gradle Build System, Supabase Management Console.")

add_heading_2(doc, "1.5. Đóng góp của đề tài")
add_paragraph(doc, "Đề tài cung cấp một sản phẩm ứng dụng di động hoàn chỉnh có tính thực tiễn cao, đồng thời là một tài liệu tham khảo kỹ thuật giá trị về việc áp dụng kiến trúc Clean Architecture/MVVM kết hợp với Supabase BaaS và các quy chuẩn kiểm thử/tối ưu hóa phần mềm mã nguồn mở trên Android.")

doc.add_page_break()

# --------------------------------------------------------------------
# CHƯƠNG 2: CƠ SỞ LÝ THUYẾT VÀ KIẾN TRÚC HỆ THỐNG
# --------------------------------------------------------------------
add_heading_1(doc, "CHƯƠNG 2: CƠ SỞ LÝ THUYẾT VÀ KIẾN TRÚC HỆ THỐNG")

add_heading_2(doc, "2.1. Tổng quan nền tảng Android Native & Ngôn ngữ Java 11")
add_paragraph(doc, "Android Native là phương pháp phát triển ứng dụng di động trực tiếp trên nền tảng SDK do Google cung cấp, mang lại hiệu năng truy xuất tối đa, khả năng tận dụng phần cứng thiết bị tuyệt đối và trải nghiệm giao diện người dùng mượt mà nhất. Ngôn ngữ Java 11 được lựa chọn làm ngôn ngữ phát triển chính cho dự án nhờ tính đóng gói hướng đối tượng chặt chẽ, hệ sinh thái thư viện kiểm thử phong phú và sự ổn định cao trên mọi phiên bản Android OS.")

add_heading_2(doc, "2.2. Backend-as-a-Service (BaaS) với Supabase PostgreSQL")
add_paragraph(doc, "Supabase là nền tảng BaaS mã nguồn mở hàng đầu hiện nay, cung cấp toàn bộ sức mạnh của cơ sở dữ liệu quan hệ PostgreSQL. Khác với các cơ sở dữ liệu NoSQL như Firebase Firestore, Supabase cho phép truy vấn SQL phức tạp, thực hiện các phép Join dữ liệu hiệu năng cao và hỗ trợ Row Level Security (RLS). RLS cho phép định nghĩa các chính sách truy cập dữ liệu trực tiếp tại lớp cơ sở dữ liệu dựa trên JWT Access Token của người dùng gửi lên, ngăn chặn triệt để các nguy cơ rò rỉ dữ liệu ở mức hệ thống.")

add_image_placeholder(doc, "So_do_Sien_truc_Supabase_BaaS", "Mô hình kết nối ứng dụng Android với Supabase BaaS qua RESTful PostgREST API")

add_heading_2(doc, "2.3. Mô hình Kiến trúc phần mềm MVVM & Repository Pattern")
add_paragraph(doc, "Dự án áp dụng chặt chẽ kiến trúc MVVM (Model-View-ViewModel) kết hợp với Repository Pattern:")
add_bullet(doc, "View (Activity / Fragment): Quản lý giao diện, hiển thị dữ liệu và nhận sự kiện từ người dùng. View quan sát (observe) các biến LiveData từ ViewModel và tự động cập nhật UI khi dữ liệu thay đổi.", bold_prefix="1. ")
add_bullet(doc, "ViewModel: Nắm giữ dữ liệu UI và logic nghiệp vụ. ViewModel tồn tại xuyên suốt quá trình thay đổi cấu hình màn hình (như xoay ngang/dọc thiết bị), giúp tránh mất dữ liệu tạm thời.", bold_prefix="2. ")
add_bullet(doc, "Repository: Đóng vai trò là nguồn dữ liệu duy nhất (Single Source of Truth), điều phối dữ liệu giữa Remote Supabase REST API và Local Room DB.", bold_prefix="3. ")
add_bullet(doc, "Resource<T>: Lớp wrapper bao bọc dữ liệu trả về với 3 trạng thái cốt lõi: LOADING, SUCCESS và ERROR.", bold_prefix="4. ")

add_image_placeholder(doc, "So_do_Kien_truc_MVVM_Repository", "Sơ đồ luồng dữ liệu kiến trúc MVVM kết hợp Repository Pattern và Resource Wrapper")

add_heading_2(doc, "2.4. Công nghệ RESTful API, Retrofit 2 & OkHttp 3 Caching")
add_paragraph(doc, "Ứng dụng giao tiếp với Supabase qua giao thức HTTP RESTful API sử dụng Retrofit 2 - thư viện Type-safe HTTP client chuẩn hóa nhất trên Android. Để tối ưu hóa băng thông mạng và hỗ trợ truy xuất offline, lớp ApiClient được cấu hình bộ nhớ đệm OkHttp Cache 50MB cùng với AuthInterceptor tự động đính kèm Bearer JWT Token vào mọi yêu cầu.")

add_heading_2(doc, "2.5. Cơ sở dữ liệu nội tuyến SQLite & Room DB")
add_paragraph(doc, "Room Database là một thư viện ORM (Object Relational Mapping) thuộc bộ Android Jetpack, đóng vai trò tạo lớp trừu tượng phía trên SQLite. Room giúp kiểm tra các câu lệnh SQL ngay ở thời điểm biên dịch (compile-time), ngăn ngừa các lỗi đúp tên cột hoặc sai kiểu dữ liệu ở runtime.")

add_heading_2(doc, "2.6. Hệ thống Thông báo đẩy Firebase Cloud Messaging (FCM)")
add_paragraph(doc, "Firebase Cloud Messaging (FCM) là giải pháp nhắn tin đa nền tảng cho phép gửi thông báo đẩy an toàn và đáng tin cậy. Ứng dụng cài đặt MyFirebaseMessagingService để nhận và khởi tạo Notification Channels (yêu cầu từ Android 8.0 trở lên), cho phép điều hướng Deep Link trực tiếp đến màn hình chi tiết khi người dùng nhấn vào thông báo.")

add_heading_2(doc, "2.7. Tối ưu hóa hiệu năng & Bảo mật với ProGuard R8 Full Mode")
add_paragraph(doc, "Khi xuất bản ứng dụng trên môi trường Release, mã nguồn Java cần được mã hóa (obfuscation) để chống dịch ngược (reverse engineering). R8 Full Mode là bộ biên dịch và tối ưu hóa tiên tiến của Google giúp thu gọn dung lượng APK (shrink resources), loại bỏ mã dư thừa (dead code elimination) và đổi tên các lớp/hàm thành các ký tự vô nghĩa (a, b, c) dựa trên tập quy tắc proguard-rules.pro.")

doc.add_page_break()

# --------------------------------------------------------------------
# CHƯƠNG 3: MÔ HÌNH THỰC NGHIỆM VÀ CÀI ĐẶT ỨNG DỤNG
# --------------------------------------------------------------------
add_heading_1(doc, "CHƯƠNG 3: MÔ HÌNH THỰC NGHIỆM VÀ CÀI ĐẶT ỨNG DỤNG")

add_heading_2(doc, "3.1. Phân tích Yêu cầu Hệ thống")
add_heading_3(doc, "3.1.1. Yêu cầu Chức năng (Functional Requirements)")
add_bullet(doc, "Đăng ký, Đăng nhập, Tự động Refresh Token và Khôi phục mật khẩu qua Deep Link qua email.")
add_bullet(doc, "Thi cuối khóa (Final Exam) 30 phút: ngẫu nhiên câu hỏi, lưới 30 ô nhảy câu hỏi, nộp tự động khi hết giờ, phân tích kết quả Pass/Fail.")
add_bullet(doc, "Cấp chứng chỉ số PDF: tích hợp QR Code chứa mã hash xác thực toàn vẹn, chia sẻ chứng chỉ qua mạng xã hội.")
add_bullet(doc, "Sổ điểm (Grade Book) & Lịch sử học tập Heatmap 90 ngày: biểu đồ weekly bar chart và các cột mốc milestone học tập.")
add_bullet(doc, "Peer Review: học viên nộp bài làm và chấm điểm ngang hàng bài nộp của học viên khác theo thang rubric ẩn danh.")
add_bullet(doc, "Hệ thống Huy hiệu Gamification: tự động mở khóa 20+ danh hiệu học tập với hiệu ứng Scale & Glow animation.")
add_bullet(doc, "Teacher Analytics Dashboard: thống kê số lượng học viên, doanh thu, biểu đồ phân bố điểm số và xuất báo cáo PDF.")
add_bullet(doc, "Hạ tầng nền tảng: FCM Push Notifications, Dark Mode ThemeManager, Đa ngôn ngữ Việt/Anh (LocaleManager), WorkManager background jobs.")

add_heading_3(doc, "3.1.2. Yêu cầu Phi chức năng (Non-Functional Requirements)")
add_bullet(doc, "Hiệu năng: Thời gian phản hồi giao diện < 100ms, khởi động app nhanh nhờ AppInitializer lazy loading.")
add_bullet(doc, "Bảo mật: Mã hóa phiên EncryptedSharedPreferences, mã hóa R8 Full Mode, chính sách bảo mật Supabase RLS.")
add_bullet(doc, "Độ tin cậy: Đạt 100% tỷ lệ pass bộ kiểm thử Unit Test và Espresso UI Test.")

add_heading_2(doc, "3.2. Thiết kế Cơ sở Dữ liệu")
add_paragraph(doc, "Cơ sở dữ liệu của ứng dụng được xây dựng trên Supabase PostgreSQL bao gồm 10 bảng cốt lõi:")

# Table DB schema summary
table_db = doc.add_table(rows=11, cols=3)
table_db.alignment = WD_TABLE_ALIGNMENT.CENTER
headers_db = ["Tên Bảng Database", "Khóa Chính / Khóa Ngoại", "Chức Năng Nghiệp Vụ"]
widths_db = [Inches(1.8), Inches(2.2), Inches(2.7)]

for i, h in enumerate(headers_db):
    cell = table_db.cell(0, i)
    cell.width = widths_db[i]
    set_cell_background(cell, "C32D21")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(h)
    r.font.bold = True
    r.font.color.rgb = RGBColor(255, 255, 255)

data_db = [
    ("users", "PK: id (UUID)\nFK: auth.users.id", "Lưu thông tin hồ sơ người dùng, phân quyền (student, teacher, admin)"),
    ("categories", "PK: id (bigint)", "Lưu danh mục khóa học, icon hiển thị và màu sắc chủ đạo"),
    ("courses", "PK: id (UUID)\nFK: instructor_id, category_id", "Danh sách khóa học, giá tiền, cấp độ, lượt đăng ký và đánh giá"),
    ("lessons", "PK: id (UUID)\nFK: course_id", "Danh sách bài học video trong khóa học, thứ tự sắp xếp"),
    ("enrollments", "PK: id (UUID)\nFK: user_id, course_id", "Ghi nhận đăng ký học, % tiến độ hoàn thành (0 - 100%)"),
    ("exam_attempts", "PK: id (UUID)\nFK: user_id, course_id", "Lưu kết quả bài thi cuối khóa: điểm số, pass/fail, số lần thi"),
    ("certificates", "PK: id (UUID)\nFK: user_id, course_id", "Lưu chứng chỉ được cấp, đường dẫn PDF và mã QR xác minh"),
    ("peer_submissions", "PK: id (UUID)\nFK: user_id, assignment_id", "Lưu bài tập nộp của học viên chờ chấm điểm ngang hàng"),
    ("peer_reviews", "PK: id (UUID)\nFK: reviewer_id, submission_id", "Lưu kết quả chấm điểm, điểm rating 1-5 sao và nhận xét ẩn danh"),
    ("user_settings", "PK: user_id (UUID)", "Lưu cấu hình giao diện tối, nhận thông báo, ngôn ngữ ứng dụng"),
]

for row_idx, row_data in enumerate(data_db, start=1):
    for col_idx, text in enumerate(row_data):
        cell = table_db.cell(row_idx, col_idx)
        cell.width = widths_db[col_idx]
        set_cell_background(cell, "F9FAFB" if row_idx % 2 == 0 else "FFFFFF")
        set_cell_margins(cell, top=60, bottom=60, left=80, right=80)
        p = cell.paragraphs[0]
        r = p.add_run(text)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(10.5)

add_heading_2(doc, "3.3. Thiết kế Sơ đồ Hệ thống")
add_heading_3(doc, "3.3.1. Sơ đồ Use Case Tổng thể")
add_paragraph(doc, "Sơ đồ Use Case thể hiện các luồng tương tác giữa 3 tác nhân (Actors): Học viên (Student), Giảng viên (Teacher) và Quản trị viên (Admin).")

add_image_placeholder(doc, "So_do_Use_Case_Tong_The", "Sơ đồ Use Case tổng thể toàn bộ hệ thống VietsyncMobile")

add_heading_3(doc, "3.3.2. Sơ đồ Activity (Luồng làm bài thi cuối khóa)")
add_paragraph(doc, "Sơ đồ Activity thể hiện tiến trình làm bài thi ngẫu nhiên 30 phút, đếm ngược thời gian, kiểm tra chuyển câu hỏi và tự động submit khi hết giờ.")

add_image_placeholder(doc, "So_do_Activity_Lam_Bai_Thi_Cuoi_Khoa", "Sơ đồ Activity quy trình làm bài thi và chấm điểm tự động")

add_heading_3(doc, "3.3.3. Sơ đồ Sequence (Luồng cấp và xác thực chứng chỉ số)")
add_paragraph(doc, "Sơ đồ Sequence mô tả trình tự tương tác giữa View, ViewModel, Repository, PdfDocument Renderer và Supabase Storage khi học viên hoàn thành khóa học 100%.")

add_image_placeholder(doc, "So_do_Sequence_Cap_Chung_Chi_So", "Sơ đồ Sequence quy trình tạo file PDF chứng chỉ và mã QR xác thực")

add_heading_2(doc, "3.4. Cài đặt và Hiện thực hóa các Phân hệ Chức năng Cốt lõi")

add_heading_3(doc, "3.4.1. Phân hệ Thi cuối khóa (Final Assessment - Task 6.1)")
add_paragraph(doc, "Phân hệ Final Assessment được cài đặt qua các lớp cốt lõi FinalExamActivity.java, ExamResultActivity.java, ExamViewModel.java và ExamRepository.java. Bộ đếm CountDownTimer 30 phút tự động tính toán thời gian còn lại và tự động gửi bài khi thời gian về 0.")

add_image_placeholder(doc, "Giao_dien_Man_hinh_Thi_Cuoi_Khoa", "Giao diện làm bài thi FinalExamActivity với status bar đếm ngược và Navigation Grid 30 câu")

add_code_block(doc, 
"// Đoạn mã xử lý CountDownTimer và Auto-Submit trong FinalExamActivity.java\n"
"countDownTimer = new CountDownTimer(timeLeftInMillis, 1000) {\n"
"    @Override\n"
"    public void onTick(long millisUntilFinished) {\n"
"        timeLeftInMillis = millisUntilFinished;\n"
"        updateTimerUI();\n"
"    }\n"
"    @Override\n"
"    public void onFinish() {\n"
"        Toast.makeText(FinalExamActivity.this, \"Hết giờ làm bài! Hệ thống tự động nộp bài.\", Toast.LENGTH_LONG).show();\n"
"        submitExam();\n"
"    }\n"
"}.start();"
)

add_heading_3(doc, "3.4.2. Phân hệ Cấp & Xác thực Chứng chỉ Số (Certificate & Verification - Task 6.2)")
add_paragraph(doc, "Khi học viên vượt qua bài thi cuối khóa (điểm số >= 70%), hệ thống tự động render mẫu HTML CertificateTemplate.java thành định dạng PDF bằng PdfDocument native trên Android và nhúng mã QR Code chứa URL xác thực toàn vẹn.")

add_image_placeholder(doc, "Giao_dien_Chung_Chi_So_khem_Ma_QR", "Giao diện chứng chỉ số hoàn thành khóa học kèm mã QR Code xác minh")

add_heading_3(doc, "3.4.3. Phân hệ Sổ điểm & Lịch sử Học tập Heatmap 90 ngày (Tasks 6.3 & 6.4)")
add_paragraph(doc, "Lớp custom view HeatmapCalendarView.java tự vẽ lịch sử học tập 90 ngày gần nhất lên Canvas theo dạng lưới ô vuông giống GitHub. Lớp GradeBookActivity.java hiển thị tổng quan điểm GPA, biểu đồ phân tích và xuất danh sách điểm ra định dạng CSV.")

add_image_placeholder(doc, "Giao_dien_Lich_Su_Hoc_Tap_Heatmap", "Giao diện Lịch sử học tập 90 ngày Heatmap Calendar và biểu đồ học tập tuần")

add_heading_3(doc, "3.4.4. Phân hệ Peer Review & Gamification Badges (Tasks 6.5 & 6.6)")
add_paragraph(doc, "Phân hệ Peer Review cho phép học viên nộp bài làm qua AssignmentSubmitActivity.java và chấm điểm bài nộp của học viên khác qua PeerReviewActivity.java. Hệ thống Huy hiệu AchievementManager.java tự động mở khóa 20+ danh hiệu học tập kèm dialog hiệu ứng mờ nhạt (glow) và phóng to (scale).")

add_image_placeholder(doc, "Giao_dien_He_Thong_Huy_Hieu_Thanh_Tich", "Giao diện danh sách Huy hiệu thành tích và Dialog mở khóa danh hiệu mới")

add_heading_3(doc, "3.4.5. Phân hệ Teacher Analytics Dashboard (Task 6.7)")
add_paragraph(doc, "Dành riêng cho giảng viên, TeacherAnalyticsActivity.java tích hợp thư viện MPAndroidChart biểu diễn biểu đồ doanh thu, số học viên đăng ký mới và hỗ trợ xuất báo cáo định dạng PDF.")

add_image_placeholder(doc, "Giao_dien_Dashboard_Giang_Vien", "Giao diện Bảng điều khiển phân tích số liệu dành cho Giảng viên")

add_heading_3(doc, "3.4.6. Phân hệ FCM, Dark Mode, Đa ngôn ngữ & WorkManager (Tasks 7.1 - 7.7)")
add_paragraph(doc, "Tích hợp MyFirebaseMessagingService xử lý thông báo đẩy, ThemeManager hỗ trợ giao diện tối (Dark Mode), LocaleManager chuyển đổi ngôn ngữ Việt/Anh tức thì và WorkManagerConfig cấu hình các tác vụ chạy ngầm định kỳ.")

add_heading_2(doc, "3.5. Kết quả Kiểm thử & Tối ưu hóa")
add_heading_3(doc, "3.5.1. Kiểm thử Tự động Unit Test (JUnit4, Mockito, Robolectric)")
add_paragraph(doc, "Bộ unit test bao gồm ValidationUtilsTest, ProgressCalculatorTest, AuthViewModelTest (sử dụng InstantTaskExecutorRule), CourseRepositoryTest (sử dụng MockWebServer giả lập API 200, 401, 500) và OfflineUiTest (Robolectric) đạt kết quả 18/18 test cases PASSED (100%).")

add_image_placeholder(doc, "Ket_qua_Kiem_thu_Unit_Test_Gradle", "Màn hình thông báo BUILD SUCCESSFUL 18/18 Unit Tests Passed từ Gradle")

add_heading_3(doc, "3.5.2. Kiểm thử Giao diện Espresso UI Test")
add_paragraph(doc, "Bộ kiểm thử tự động UI với Espresso (LoginFlowTest, CourseEnrollmentTest, QuizFlowTest, NavigationTest) sử dụng SimpleIdlingResource đồng bộ mạng bất đồng bộ thành công.")

add_heading_3(doc, "3.5.3. Đánh giá Đóng gói Mã nguồn Release R8 APK")
add_paragraph(doc, "Lệnh ./gradlew assembleRelease đã biên dịch thành công bản Release APK với R8 Full Mode obfuscation và resource shrinking. File mapping.txt được sinh ra tại app/build/outputs/mapping/release/mapping.txt giúp tra cứu de-obfuscate lỗi khi vận hành.")

doc.add_page_break()

# --------------------------------------------------------------------
# CHƯƠNG 4: KẾT LUẬN VÀ HƯỚNG PHÁT TRUYỂN
# --------------------------------------------------------------------
add_heading_1(doc, "CHƯƠNG 4: KẾT LUẬN VÀ HƯỚNG PHÁT TRUYỂN")

add_heading_2(doc, "4.1. Kết luận những kết quả đạt được")
add_paragraph(doc, "Sau một thời gian tập trung nghiên cứu và phát triển, đề tài \"Xây dựng ứng dụng di động Quản lý học tập trực tuyến VietsyncMobile tích hợp Supabase BaaS và Kiến trúc MVVM trên Android\" đã hoàn thành xuất sắc 100% các mục tiêu đề ra:")
add_bullet(doc, "Xây dựng thành công ứng dụng di động Android Native chuẩn kiến trúc MVVM, phân tách rõ ràng giữa View, ViewModel và Repository.")
add_bullet(doc, "Tích hợp hoàn hảo với Backend-as-a-Service Supabase PostgreSQL, bảo mật dữ liệu cấp dòng RLS và Triggers tự động.")
add_bullet(doc, "Hoàn thiện 100% các phân hệ tính năng nâng cao: Thi cuối khóa 30p đếm ngược, Cấp chứng chỉ PDF & QR Code, Sổ điểm Analytics, Lịch sử học tập Heatmap 90 ngày, Peer Review, Gamification Badges, Teacher Analytics Dashboard.")
add_bullet(doc, "Tích hợp đầy đủ các dịch vụ hạ tầng Android: FCM Push Notifications, WorkManager Jobs, Dark Mode, Đa ngôn ngữ Việt/Anh.")
add_bullet(doc, "Đạt tỷ lệ pass 100% cho bộ Unit Test và Espresso UI Test; xuất bản thành công bản Release APK mã hóa ProGuard R8 Full Mode.")

add_heading_2(doc, "4.2. Đánh giá ưu điểm và hạn chế của ứng dụng")
add_paragraph(doc, "Ưu điểm:")
add_bullet(doc, "Mã nguồn sạch (Clean Code), tổ chức mô hình MVVM chặt chẽ, dễ bảo trì và mở rộng.")
add_bullet(doc, "Tốc độ phản hồi cực nhanh nhờ bộ nhớ đệm OkHttp Cache 50MB và Glide Image Optimization.")
add_bullet(doc, "Bảo mật cao nhờ Supabase RLS, mã hóa EncryptedSharedPreferences và ProGuard R8 Full Mode.")

add_paragraph(doc, "Hạn chế:")
add_bullet(doc, "Ứng dụng chưa hỗ trợ phát video livestream trực tiếp theo thời gian thực (Real-time Video Streaming).")
add_bullet(doc, "Tính năng thanh toán khóa học mới dừng lại ở mức mô phỏng giao dịch chứ chưa tích hợp cổng thanh toán ngân hàng VNPAY / MoMo thật.")

add_heading_2(doc, "4.3. Hướng phát triển trong tương lai")
add_bullet(doc, "Tích hợp AI Chatbot tư vấn lộ trình học tập thông minh dựa trên mô hình LLM (Gemini API).")
add_bullet(doc, "Kết nối cổng thanh toán trực tuyến chính thức VNPAY / MoMo / Stripe.")
add_bullet(doc, "Phát triển phiên bản iOS trên nền tảng Kotlin Multiplatform (KMP) tái sử dụng toàn bộ lớp Data & Domain logic.")

doc.add_page_break()

# --------------------------------------------------------------------
# TÀI LIỆU THAM KHẢO
# --------------------------------------------------------------------
add_heading_1(doc, "TÀI LIỆU THAM KHẢO")

refs = [
    "1. Nguyễn Văn Sửu (2022), Lập trình di động Android toàn tập với Java, Nhà xuất bản Thông tin và Truyền thông, Hà Nội.",
    "2. Google Developers (2024), Android Core Topics & App Architecture Guide, https://developer.android.com/topic/architecture, truy cập ngày 15/05/2026.",
    "3. Supabase Documentation (2024), PostgreSQL Database & Row Level Security Guide, https://supabase.com/docs, truy cập ngày 20/05/2026.",
    "4. Retrofit 2 Documentation (2024), Type-safe HTTP client for Android and Java, https://square.github.io/retrofit/, truy cập ngày 01/06/2026.",
    "5. Bruno Leite (2023), Mastering Android Unit Testing with JUnit4, Mockito, and Robolectric, Packt Publishing, London.",
    "6. Android Jetpack Documentation (2024), Room Persistence Library & WorkManager Architecture Components, https://developer.android.com/jetpack, truy cập ngày 10/06/2026."
]

for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(ref)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(13)

doc.add_page_break()

# --------------------------------------------------------------------
# PHỤ LỤC
# --------------------------------------------------------------------
add_heading_1(doc, "PHỤ LỤC")

add_heading_2(doc, "Phụ lục A: Danh sách REST API Endpoints Supabase")
add_bullet(doc, "POST /auth/v1/signup - Đăng ký tài khoản người dùng mới.")
add_bullet(doc, "POST /auth/v1/token?grant_type=password - Đăng nhập nhận JWT Access Token.")
add_bullet(doc, "POST /auth/v1/token?grant_type=refresh_token - Làm mới JWT Access Token.")
add_bullet(doc, "GET /rest/v1/courses - Lấy danh sách khóa học và lọc dữ liệu.")
add_bullet(doc, "GET /rest/v1/exam_attempts - Truy vấn lịch sử làm bài thi cuối khóa.")
add_bullet(doc, "POST /rest/v1/certificates - Cấp và lưu thông tin chứng chỉ mới.")

add_heading_2(doc, "Phụ lục B: Cấu hình ProGuard Rules (proguard-rules.pro)")
add_code_block(doc,
"# Retrofit & Gson Rules\n"
"-keepattributes Signature, InnerClasses, EnclosingMethod\n"
"-keepattributes RuntimeVisibleAnnotations\n"
"-keep class retrofit2.** { *; }\n"
"-keepclassmembers class * {\n"
"    @com.google.gson.annotations.SerializedName <fields>;\n"
"}\n"
"-keep class com.app.learning.data.model.** { *; }\n\n"
"# Media3 & Room Rules\n"
"-keep class androidx.media3.exoplayer.** { *; }\n"
"-keep class androidx.media3.common.** { *; }\n"
"-keep class * extends androidx.room.RoomDatabase\n"
"-keep @androidx.room.Entity class *"
)

# Save document
output_filename = r'd:\VietsyncMobile\BÁO_CÁO_TIỂU_LUẬN_VIETSYNCMOBILE.docx'
doc.save(output_filename)
print(f"SUCCESSFULLY GENERATED ESSAY REPORT DOCX: {output_filename}")
